from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import Dict, List
import base64
import re
import asyncio
from io import BytesIO 
from datetime import datetime
import extract_msg
import os
from dotenv import load_dotenv
import concurrent.futures
import pandas as pd
from docx import Document
from openai import AsyncAzureOpenAI
from starlette.responses import StreamingResponse 
from fastapi.responses import JSONResponse 
import json 
from docx.shared import Inches

# Load environment variables
if os.getenv("FLASK_ENV") != "production":
    load_dotenv()

api_key = os.getenv("AZURE_OPENAI_API_key")
azure_endpoint = os.getenv("AZURE_OPENAI")
model_name = "gpt-4o-mini"
if not api_key or not azure_endpoint:
    raise EnvironmentError("AZURE_OPENAI_API_key or AZURE_OPENAI not set!")

# Initialize AsyncAzureOpenAI client
client = AsyncAzureOpenAI(
    api_key=api_key,
    azure_endpoint=azure_endpoint,
    api_version="2024-08-01-preview",
)

# FastAPI app initialization
app = FastAPI()

# Semaphore to limit concurrent API requests
semaphore = asyncio.Semaphore(5)

# Request models
class EmailRequest(BaseModel):
    content: str  # Base64-encoded .msg file

class ExcelRequest(BaseModel):
    info: Dict[str, str] 
    
# Define request model
class WordRequest(BaseModel):
    data: str  # JSON string
    document: str  # Base64 encoded Word file

class Prompt(BaseModel):
    input: str

# Extract information from .msg file
def sync_extract_msg(file_path):
    """Extracts metadata and body from a .msg file."""
    msg = extract_msg.Message(file_path)
    email_date = msg.date
    body = re.sub(r'<[^>]+>', '', msg.body)  # Remove HTML tags
    body = re.sub(r'\s+', ' ', body).strip()  # Normalize whitespace
    return email_date, body



# Generate Stream
async def stream_processor(response):
    async for chunk in response:
        if len(chunk.choices) > 0:
            delta = chunk.choices[0].delta
            if delta.content:
                yield delta.content 
                
async def fetch_all_info(body):
    """Fetches all required information and summaries in a single prompt using Azure OpenAI."""
    # Single prompt to extract all information and summarize specific fields
    prompt = """
    Extract the following information from the email body and return it in JSON format with the keys specified below:
    - Project Title: Provide the project title only (e.g., Real Estate Fund Portfolio Management).
    - Client Name: Provide the client name (not Lionpoint) only (e.g., Bain Capital LP).
    - Use Case: Provide the specific use cases only (e.g., Real Estate Fund Forecasting, Waterfall, Asset Management, Workforce Planning).
    - Completion Date: Provide the completion date (Month and Year only, e.g., December 2021).
    - Project Objectives: Extract the main objectives of the project.
    - Business Challenges: Extract the key business challenges faced by the client.
    - Our Approach: Extract the approach taken during the project.
    - Value Created: Extract the value created or outcomes achieved from the project.
    - Measures of Success: Extract the measures of success for the project.
    - Industry: Extract the industry related to the project.

    Additionally, provide a brief summary for the following fields to use for excel database:
    - Project Objectives: Summarize the main objectives briefly listing keywords.
    - Business Challenges: Summarize the key business challenges briefly listing keyword.
    - Our Approach: Summarize the approach taken during the project briefly listing key words.
    - Value Created: Summarize the value created or outcomes achieved briefly listing keywords.
    - Measures of Success: Summarize the measures of success briefly listing keywords. 
    - Internal Resources: List the employees who worked on this project.

    Return the response in the following JSON format:
    {
        "extracted_info": {
            "Project Title": "<project_title>",
            "Client Name": "<client_name>",
            "Use Case": "<use_case>",
            "Completion Date": "<completion_date>",
            "Project Objectives": "<project_objectives>",
            "Business Challenges": "<business_challenges>",
            "Our Approach": "<our_approach>",
            "Value Created": "<value_created>",
            "Measures of Success": "<measures_of_success>",
            "Industry": "<industry>"
        },
        "summarized_info": {
            "Project Objectives": "<summary_of_project_objectives>",
            "Business Challenges": "<summary_of_business_challenges>",
            "Our Approach": "<summary_of_our_approach>",
            "Value Created": "<summary_of_value_created>",
            "Measures of Success": "<summary_of_measures_of_success>",
            "Internal Resources": "<employees>"
        }
    }
    """

    payload = {
        "messages": [
            {"role": "system", "content": "You are a helpful assistant that extracts structured information from emails and returns it in JSON format."},
            {"role": "user", "content": f"{prompt}\n\nEmail Body:\n{body}"}
        ]
    }

    try:
        async with semaphore:
            response = await client.chat.completions.create(
                model=model_name,
                messages=payload["messages"],
                response_format={"type": "json_object"}  # Ensure the response is in JSON format
            )
            # Parse the JSON response
            return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {str(e)}"

async def process_email_content(file_path):
    """Extracts and summarizes information from the .msg file."""
    loop = asyncio.get_event_loop()
    with concurrent.futures.ThreadPoolExecutor() as pool:
        email_date, body = await loop.run_in_executor(pool, sync_extract_msg, file_path)

    # Fetch all information and summaries in a single prompt
    info_json = await fetch_all_info(body)

    # Parse the JSON response
    try:
        info = json.loads(info_json)
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="Failed to parse the response from OpenAI.")

    # Handle completion date fallback
    if info["extracted_info"]["Completion Date"] == "Not Provided" or not re.search(
        r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December) \d{4}\b',
        info["extracted_info"]["Completion Date"],
        re.IGNORECASE
    ):
        if email_date:
            info["extracted_info"]["Completion Date"] = email_date.strftime("%B %Y")

    return info

async def add_heading_and_text(doc, heading, text, style=None):
    # Add the section heading
    doc.add_heading(heading, level=2)

    if isinstance(text, str) and text.strip():  # Check if text is non-empty
        paragraphs = text.split('\n')  # Split by new line
        numbered = False  # Flag to track if inside a numbered list
        last_numbered_paragraph = None

        for line in paragraphs:
            line = line.strip()  # Remove whitespace
            if not line:
                continue  # Skip empty lines

            segments = line.split('**')  # Split for bold text
            paragraph = doc.add_paragraph()  # Create a new paragraph

            if len(line) > 0 and line[0].isdigit() and ('.' in line[:3] or ')' in line[:3]):  # Check for numbered list
                numbered = True
                last_numbered_paragraph = paragraph
                number_text = line.split(maxsplit=1)  # Separate number from text
                paragraph.add_run(number_text[0] + ' ').bold = True  # Bold the number
                if len(number_text) > 1:
                    for i, segment in enumerate(number_text[1].split('**')):
                        run = paragraph.add_run((' ' if i % 2 == 1 else '') + segment.strip() + (' ' if i % 2 == 1 else ''))
                        if i % 2 == 1:
                            run.bold = True  # Bold text
                paragraph.paragraph_format.left_indent = Inches(0.50)
            elif len(line) > 0 and (line.startswith('•') or line.startswith('-')):  # Check for bullet point
                bullet_text = line.lstrip('• -').strip()
                bullet_paragraph = last_numbered_paragraph if numbered else doc.add_paragraph()
                for i, segment in enumerate(bullet_text.split('**')):
                    run = bullet_paragraph.add_run((' ' if i % 2 == 1 else '') + segment.strip() + (' ' if i % 2 == 1 else ''))
                    if i % 2 == 1:
                        run.bold = True  # Bold text
                bullet_paragraph.style = 'List Bullet'
                bullet_paragraph.paragraph_format.left_indent = Inches(1.0 if numbered else 0.50)  # Indent appropriately
            else:  # Regular text
                numbered = False  # Reset numbered flag if normal text appears
                for i, segment in enumerate(segments):
                    run = paragraph.add_run((' ' if i % 2 == 1 else '') + segment.strip() + (' ' if i % 2 == 1 else ''))  # Add spaces around bold text
                    if i % 2 == 1:
                        run.bold = True  # Bold text
    
    elif text:
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
async def create_summary_doc(existing_doc_bytes, all_data):
    # Decode the base64 document
    doc_content = base64.b64decode(existing_doc_bytes)

    # Create a BytesIO object to hold the document in memory
    doc_stream = BytesIO(doc_content)

    # Open the document
    doc = Document(doc_stream)

    # Process the provided project data
    for project in all_data:
        # Adding project title
        doc.add_heading(project['Project Title'], level=1)

        # Using the add_heading_and_text utility to add client information
        await add_heading_and_text(doc, 'Client Name:', project['Client Name'], 'Body Text')
        await add_heading_and_text(doc, 'Use Case:', project['Use Case'], 'Body Text')
        await add_heading_and_text(doc, 'Industry:', project['Industry'], 'Body Text')
        await add_heading_and_text(doc, 'Completion Date:', project['Completion Date'], 'Body Text')

        # Adding sections to the document
        await add_heading_and_text(doc, 'Project Objectives:', project['Project Objectives'], None)
        await add_heading_and_text(doc, 'Business Challenges:', project['Business Challenges'], None)
        await add_heading_and_text(doc, 'Our Approach:', project['Our Approach'], None)
        await add_heading_and_text(doc, 'Value Created:', project['Value Created'], None)
        await add_heading_and_text(doc, 'Measures of Success:', project['Measures of Success'], None)

    # Save the modified document to bytes
    updated_doc_bytes = BytesIO()
    doc.save(updated_doc_bytes)

    # Seek to the beginning to prepare for reading
    updated_doc_bytes.seek(0)

    # Convert the updated document to base64 for response
    encoded_doc = base64.b64encode(updated_doc_bytes.getvalue()).decode('utf-8')
    return encoded_doc

# API Endpoint for streaming
@app.post("/stream")
async def stream(prompt: Prompt):
    azure_open_ai_response = await client.chat.completions.create(
        model = model_name,
        messages=[{"role": "user", "content": prompt.input}],
        stream=True
    )

    return StreamingResponse(stream_processor(azure_open_ai_response), media_type="text/event-stream")

# API Endpoint for processing email
@app.post("/process-email")
async def process_email(request: EmailRequest):
    try:
        base64_content = request.content
        if not base64_content:
            raise HTTPException(status_code=400, detail="No content provided.")

        # Decode the .msg file
        decoded_content = base64.b64decode(base64_content)
        email_stream = BytesIO(decoded_content)

        # Process the email content
        result = await process_email_content(email_stream)


        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# API Endpoint for Word documentation
@app.post("/word")
async def word_documentation(request: WordRequest):
    try:
        # Extract document and project data
        document_base64 = request.document
        project_data_string = request.data
        
        # Parse the JSON string from 'data' into a list of dictionaries
        project_data = json.loads(project_data_string)
        
        if not document_base64 or not project_data:
            return JSONResponse(content={'error': 'Missing document or data'}, status_code=400)

        # Generate the updated document
        updated_doc_base64 = await create_summary_doc(document_base64, project_data)
        
        # Return the updated document as base64
        return JSONResponse(content={'modifiedDocument': updated_doc_base64}, status_code=200)
    except Exception as e:
        return JSONResponse(content={'error': str(e)}, status_code=500)
        
# API Endpoint for Excel documentation
@app.post("/excel")
async def excel_documentation(request: ExcelRequest):
    try:
        data = request.info
        df = pd.DataFrame([data])

        modified_excel_bytes = BytesIO()
        df.to_excel(modified_excel_bytes, index=False)

        return {"excel": base64.b64encode(modified_excel_bytes.getvalue()).decode('utf-8')}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

